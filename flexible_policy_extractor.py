"""
Flexible Policy Document Extractor
Adapts to varying policy document structures (not just numbered sections).
Supports PDF and Word (.docx). Same backend is used for both.

Backends:
- pdfplumber (default): page-level text + heuristic headings; can truncate nested structure.
- docling: layout-aware parsing for PDF and DOCX; preserves sections/tables/lists (pip install docling).
- unstructured: element-level partition for PDF and DOCX (pip install 'unstructured[pdf]' and/or 'unstructured[docx]').
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    partition_pdf = None
    partition_docx = None
    UNSTRUCTURED_AVAILABLE = False


@dataclass
class PolicyPassage:
    """Represents a passage from a policy document"""
    policy_id: str
    policy_name: str
    passage_id: str  # Auto-generated sequential ID
    passage_text: str
    heading: Optional[str] = None  # Section heading if available
    page_number: int = 0
    metadata: Dict = None


class FlexiblePolicyExtractor:
    """
    Flexible policy extractor that works with varying document structures.
    Use backend="docling" or "unstructured" to preserve nested structure (sections, tables, lists).
    """

    def __init__(
        self,
        min_passage_length: int = 100,
        max_passage_length: int = 1000,
        backend: str = "pdfplumber",
    ):
        self.min_passage_length = min_passage_length
        self.max_passage_length = max_passage_length
        self.backend = (backend or "pdfplumber").lower().strip()
        if self.backend not in ("pdfplumber", "docling", "unstructured"):
            raise ValueError(f"backend must be pdfplumber, docling, or unstructured; got {backend!r}")

    def extract_from_pdf(self, pdf_path: str) -> List[PolicyPassage]:
        """Extract passages from PDF. Uses backend: pdfplumber (default), docling, or unstructured."""
        if self.backend == "docling":
            return self._extract_from_pdf_docling(pdf_path)
        if self.backend == "unstructured":
            return self._extract_from_pdf_unstructured(pdf_path)
        return self._extract_from_pdf_pdfplumber(pdf_path)

    def _extract_from_pdf_pdfplumber(self, pdf_path: str) -> List[PolicyPassage]:
        """Extract passages from PDF using pdfplumber (page text + heuristic headings)."""
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber required")
        
        print(f"Extracting from PDF: {pdf_path}")
        
        pdf_file = Path(pdf_path)
        policy_id = pdf_file.stem
        policy_name = self._extract_policy_name_from_filename(pdf_file.name)
        
        passages = []
        current_heading = None
        current_text = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                if page_num % 5 == 0:
                    print(f"  Processing page {page_num}/{len(pdf.pages)}...", end='\r')
                
                # Extract text with formatting info
                text = page.extract_text()
                if not text:
                    continue
                
                # Try to extract structured content
                words = page.extract_words()
                
                # Identify potential headings (short lines, possibly bold/larger)
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if this looks like a heading
                    is_heading = self._is_heading(line, words, page)
                    
                    if is_heading:
                        # Save previous passage
                        if current_text:
                            passage_text = ' '.join(current_text)
                            if len(passage_text) >= self.min_passage_length:
                                passages.append(self._create_passage(
                                    policy_id, policy_name, passage_text, 
                                    current_heading, page_num - 1, len(passages) + 1
                                ))
                        
                        # New heading
                        current_heading = line
                        current_text = []
                    else:
                        # Regular text
                        current_text.append(line)
            
            # Don't forget last passage
            if current_text:
                passage_text = ' '.join(current_text)
                if len(passage_text) >= self.min_passage_length:
                    passages.append(self._create_passage(
                        policy_id, policy_name, passage_text,
                        current_heading, page_num, len(passages) + 1
                    ))
        
        # If no headings found, create passages by chunking text
        if len(passages) == 0:
            print("  No headings detected. Creating passages by text chunking...")
            passages = self._chunk_by_text(pdf_path, policy_id, policy_name)

        print(f"\n✓ Extracted {len(passages)} passages from PDF (pdfplumber)")
        return passages

    def _passages_from_markdown(
        self, md: str, policy_id: str, policy_name: str, source_label: str
    ) -> List[PolicyPassage]:
        """Build passages from Docling-style markdown (shared by PDF and DOCX docling)."""
        if not md or not md.strip():
            return []
        section_blocks = re.split(r'\n(?=#{1,6}\s+)', md.strip())
        passages = []
        for block in section_blocks:
            block = block.strip()
            if not block:
                continue
            lines = block.split("\n")
            heading = None
            if lines and lines[0].startswith("#"):
                heading = re.sub(r'^#+\s*', '', lines[0]).strip()
                body = "\n".join(lines[1:]).strip()
            else:
                body = block
            text = (heading + "\n\n" + body).strip() if heading else body
            if not text or len(text) < 30:
                continue
            if len(text) >= self.min_passage_length or (heading and len(body) >= 50):
                passages.append(self._create_passage(
                    policy_id, policy_name, text, heading, 0, len(passages) + 1
                ))
        if len(passages) == 0 and md.strip():
            text = md.strip()[: self.max_passage_length * 2]
            if len(text) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, text[: self.max_passage_length], None, 0, 1
                ))
        return passages

    def _passages_from_unstructured_elements(
        self, elements: list, policy_id: str, policy_name: str, source_label: str
    ) -> List[PolicyPassage]:
        """Build passages from Unstructured elements (shared by PDF and DOCX)."""
        passages = []
        current_heading = None
        current_parts = []
        for el in elements:
            cat = getattr(el, "category", None) or ""
            text = (getattr(el, "text", None) or "").strip()
            if not text:
                continue
            if cat == "Title" or (cat == "Header" and len(text) < 150):
                if current_parts:
                    passage_text = "\n\n".join(current_parts)
                    if len(passage_text) >= self.min_passage_length:
                        passages.append(self._create_passage(
                            policy_id, policy_name, passage_text,
                            current_heading, 0, len(passages) + 1
                        ))
                current_heading = text
                current_parts = []
            else:
                current_parts.append(text)
        if current_parts:
            passage_text = "\n\n".join(current_parts)
            if len(passage_text) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, passage_text,
                    current_heading, 0, len(passages) + 1
                ))
        if len(passages) == 0 and elements:
            full = "\n\n".join(
                (getattr(e, "text", "") or "").strip() for e in elements if getattr(e, "text", None)
            )
            if len(full) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, full[: self.max_passage_length], None, 0, 1
                ))
        return passages

    def _extract_from_pdf_docling(self, pdf_path: str) -> List[PolicyPassage]:
        """Extract passages using Docling (preserves sections, tables, lists)."""
        if not DOCLING_AVAILABLE:
            raise ImportError("docling required for backend='docling'. Install: pip install docling")
        pdf_file = Path(pdf_path)
        policy_id = pdf_file.stem
        policy_name = self._extract_policy_name_from_filename(pdf_file.name)
        print(f"Extracting from PDF (Docling): {pdf_path}")
        converter = DocumentConverter()
        result = converter.convert(pdf_path)
        doc = result.document
        md = doc.export_to_markdown()
        if not md or not md.strip():
            print("  No content from Docling; falling back to empty list.")
            return []
        passages = self._passages_from_markdown(md, policy_id, policy_name, "docling")
        print(f"\n✓ Extracted {len(passages)} passages from PDF (docling)")
        return passages

    def _extract_from_pdf_unstructured(self, pdf_path: str) -> List[PolicyPassage]:
        """Extract passages using Unstructured (Title/NarrativeText/Table elements)."""
        if not UNSTRUCTURED_AVAILABLE or partition_pdf is None:
            raise ImportError(
                "unstructured required for backend='unstructured'. Install: pip install 'unstructured[pdf]'"
            )
        pdf_file = Path(pdf_path)
        policy_id = pdf_file.stem
        policy_name = self._extract_policy_name_from_filename(pdf_file.name)
        print(f"Extracting from PDF (Unstructured): {pdf_path}")
        elements = partition_pdf(pdf_path)
        passages = self._passages_from_unstructured_elements(
            elements, policy_id, policy_name, "unstructured"
        )
        print(f"\n✓ Extracted {len(passages)} passages from PDF (unstructured)")
        return passages

    def extract_from_docx(self, docx_path: str) -> List[PolicyPassage]:
        """Extract passages from DOCX. Uses backend when set to docling/unstructured; else python-docx."""
        if self.backend == "docling":
            return self._extract_from_docx_docling(docx_path)
        if self.backend == "unstructured":
            return self._extract_from_docx_unstructured(docx_path)
        return self._extract_from_docx_python_docx(docx_path)

    def _extract_from_docx_docling(self, docx_path: str) -> List[PolicyPassage]:
        """Extract from DOCX using Docling (same structure preservation as PDF)."""
        if not DOCLING_AVAILABLE:
            raise ImportError("docling required for backend='docling'. Install: pip install docling")
        docx_file = Path(docx_path)
        policy_id = docx_file.stem
        policy_name = self._extract_policy_name_from_filename(docx_file.name)
        print(f"Extracting from DOCX (Docling): {docx_path}")
        converter = DocumentConverter()
        result = converter.convert(docx_path)
        doc = result.document
        md = doc.export_to_markdown()
        if not md or not md.strip():
            return []
        passages = self._passages_from_markdown(md, policy_id, policy_name, "docling")
        print(f"\n✓ Extracted {len(passages)} passages from DOCX (docling)")
        return passages

    def _extract_from_docx_unstructured(self, docx_path: str) -> List[PolicyPassage]:
        """Extract from DOCX using Unstructured (Title/NarrativeText/Table elements)."""
        if not UNSTRUCTURED_AVAILABLE or partition_docx is None:
            raise ImportError(
                "unstructured required for backend='unstructured' on Word files. "
                "Install: pip install 'unstructured[docx]'"
            )
        docx_file = Path(docx_path)
        policy_id = docx_file.stem
        policy_name = self._extract_policy_name_from_filename(docx_file.name)
        print(f"Extracting from DOCX (Unstructured): {docx_path}")
        elements = partition_docx(docx_path)
        passages = self._passages_from_unstructured_elements(
            elements, policy_id, policy_name, "unstructured"
        )
        print(f"\n✓ Extracted {len(passages)} passages from DOCX (unstructured)")
        return passages

    def _extract_from_docx_python_docx(self, docx_path: str) -> List[PolicyPassage]:
        """Extract passages from DOCX using python-docx (paragraph structure)."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required")
        print(f"Extracting from DOCX: {docx_path}")
        
        docx_file = Path(docx_path)
        policy_id = docx_file.stem
        policy_name = self._extract_policy_name_from_filename(docx_file.name)
        
        passages = []
        current_heading = None
        current_paragraphs = []
        
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading (based on style)
            is_heading = self._is_heading_docx(para)
            
            if is_heading:
                # Save previous passage
                if current_paragraphs:
                    passage_text = ' '.join(current_paragraphs)
                    if len(passage_text) >= self.min_passage_length:
                        passages.append(self._create_passage(
                            policy_id, policy_name, passage_text,
                            current_heading, 0, len(passages) + 1
                        ))
                
                # New heading
                current_heading = text
                current_paragraphs = []
            else:
                # Regular paragraph
                current_paragraphs.append(text)
        
        # Don't forget last passage
        if current_paragraphs:
            passage_text = ' '.join(current_paragraphs)
            if len(passage_text) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, passage_text,
                    current_heading, 0, len(passages) + 1
                ))
        
        # If no headings found, create passages by chunking
        if len(passages) == 0:
            print("  No headings detected. Creating passages by text chunking...")
            passages = self._chunk_by_text_docx(docx_path, policy_id, policy_name)
        
        print(f"\n✓ Extracted {len(passages)} passages from DOCX")
        return passages
    
    def _is_heading(self, line: str, words: List, page) -> bool:
        """Check if line looks like a heading"""
        # Short lines are often headings
        if len(line) < 100 and len(line.split()) < 15:
            # Check if it's all caps (common for headings)
            if line.isupper() and len(line) > 5:
                return True
            # Check if it starts with number (section number)
            if re.match(r'^\d+[\.\)]', line):
                return True
            # Check if it's a common heading pattern
            if re.match(r'^(Section|Chapter|Part|Article)\s+\d+', line, re.IGNORECASE):
                return True
            # Check if line ends without period (headings usually don't)
            if not line.endswith('.') and len(line.split()) < 10:
                return True
        
        return False
    
    def _is_heading_docx(self, para) -> bool:
        """Check if DOCX paragraph is a heading"""
        # Check style name
        style_name = para.style.name.lower()
        if 'heading' in style_name or 'title' in style_name:
            return True
        
        # Check if paragraph is short and formatted differently
        text = para.text.strip()
        if len(text) < 100 and len(text.split()) < 15:
            # Check formatting
            if para.runs:
                # Check if bold or larger font
                for run in para.runs:
                    if run.bold or (run.font.size and run.font.size > Pt(12)):
                        return True
            
            # Check patterns
            if re.match(r'^\d+[\.\)]', text):
                return True
            if re.match(r'^(Section|Chapter|Part|Article)', text, re.IGNORECASE):
                return True
        
        return False
    
    def _chunk_by_text(self, pdf_path: str, policy_id: str, policy_name: str) -> List[PolicyPassage]:
        """Create passages by chunking text when no structure detected"""
        passages = []
        all_text = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(self._clean_text(text))
        
        full_text = ' '.join(all_text)
        
        # Split by paragraphs (double newlines)
        paragraphs = re.split(r'\n\s*\n', full_text)
        
        # Group paragraphs into passages
        current_chunk = []
        chunk_length = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < 20:
                continue
            
            # If adding this paragraph would exceed max, save current chunk
            if chunk_length + len(para) > self.max_passage_length and current_chunk:
                passage_text = ' '.join(current_chunk)
                if len(passage_text) >= self.min_passage_length:
                    passages.append(self._create_passage(
                        policy_id, policy_name, passage_text, None, 0, len(passages) + 1
                    ))
                current_chunk = []
                chunk_length = 0
            
            current_chunk.append(para)
            chunk_length += len(para)
        
        # Don't forget last chunk
        if current_chunk:
            passage_text = ' '.join(current_chunk)
            if len(passage_text) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, passage_text, None, 0, len(passages) + 1
                ))
        
        return passages
    
    def _chunk_by_text_docx(self, docx_path: str, policy_id: str, policy_name: str) -> List[PolicyPassage]:
        """Create passages by chunking DOCX text"""
        passages = []
        paragraphs = []
        
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 20:
                paragraphs.append(text)
        
        # Group paragraphs into passages
        current_chunk = []
        chunk_length = 0
        
        for para in paragraphs:
            if chunk_length + len(para) > self.max_passage_length and current_chunk:
                passage_text = ' '.join(current_chunk)
                if len(passage_text) >= self.min_passage_length:
                    passages.append(self._create_passage(
                        policy_id, policy_name, passage_text, None, 0, len(passages) + 1
                    ))
                current_chunk = []
                chunk_length = 0
            
            current_chunk.append(para)
            chunk_length += len(para)
        
        # Last chunk
        if current_chunk:
            passage_text = ' '.join(current_chunk)
            if len(passage_text) >= self.min_passage_length:
                passages.append(self._create_passage(
                    policy_id, policy_name, passage_text, None, 0, len(passages) + 1
                ))
        
        return passages
    
    def _create_passage(self, policy_id: str, policy_name: str, 
                       text: str, heading: Optional[str], page_num: int, passage_num: int = 1) -> PolicyPassage:
        """Create a policy passage"""
        passage_id = f"{policy_id}_passage_{passage_num}"
        
        return PolicyPassage(
            policy_id=policy_id,
            policy_name=policy_name,
            passage_id=passage_id,
            passage_text=text[:self.max_passage_length] if len(text) > self.max_passage_length else text,
            heading=heading,
            page_number=page_num,
            metadata={'extraction_method': 'flexible'}
        )
    
    def _extract_policy_name_from_filename(self, filename: str) -> str:
        """Extract policy name from filename"""
        # Remove extension
        name = Path(filename).stem
        # Clean up common patterns
        name = re.sub(r'clientname-IS-POL-\d+-', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s+v\d+\.\d+.*$', '', name)  # Remove version
        name = re.sub(r'\s+\d+$', '', name)  # Remove trailing numbers
        return name.strip()
    
    def _clean_text(self, text: str) -> str:
        """Clean text"""
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def convert_to_compliance_format(self, passages: List[PolicyPassage]) -> List[Dict]:
        """Convert passages to compliance mapping format"""
        formatted = []
        
        for idx, passage in enumerate(passages, 1):
            # Update passage_id with proper number
            passage.passage_id = f"{passage.policy_id}_passage_{idx}"
            
            formatted.append({
                'id': passage.passage_id,
                'name': f"{passage.policy_name} - {passage.heading or 'Passage ' + str(idx)}",
                'text': passage.passage_text,
                'section': passage.heading or f"Passage {idx}",
                'metadata': {
                    'policy_id': passage.policy_id,
                    'policy_name': passage.policy_name,
                    'heading': passage.heading,
                    'page_number': passage.page_number,
                    **passage.metadata
                }
            })
        
        return formatted
    
    def extract_from_file(self, file_path: str) -> List[PolicyPassage]:
        """Extract from file (auto-detects format)"""
        file_path_obj = Path(file_path)
        
        if file_path_obj.suffix.lower() == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_path_obj.suffix.lower() in ['.docx', '.doc']:
            return self.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_path_obj.suffix}")


def load_all_policies_from_dir(policies_dir: str = "data/02_processed/policies") -> List[Dict]:
    """
    Load all policy passages from individual *_for_mapping.json files in a directory.
    Returns a single list (same schema as one _for_mapping.json). Use when a combined
    list is needed (e.g. compliance mapping, annotate_mappings).
    """
    path = Path(policies_dir)
    if not path.is_dir():
        return []
    all_passages = []
    for f in sorted(path.glob("*_for_mapping.json")):
        try:
            with open(f, "r", encoding="utf-8") as fp:
                data = json.load(fp)
            all_passages.extend(data if isinstance(data, list) else [data])
        except Exception:
            continue
    return all_passages


def extract_all_policies_flexible(
    policies_dir: str = "data/01_raw/policies",
    output_dir: str = "data/02_processed/policies",
    backend: str = "pdfplumber",
):
    """
    Extract all policies. Use backend='docling' or 'unstructured' to preserve
    nested structure (sections, tables, lists) instead of pdfplumber.
    """
    policies_path = Path(policies_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    pdf_files = list(policies_path.glob("*.pdf"))
    docx_files = list(policies_path.glob("*.docx"))
    all_files = pdf_files + docx_files

    if not all_files:
        print(f"⚠️  No policy files found in {policies_dir}")
        return

    print(f"Found {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX file(s) | backend={backend}")
    print("=" * 60)

    extractor = FlexiblePolicyExtractor(backend=backend)
    saved = []

    for policy_file in all_files:
        print(f"\nProcessing: {policy_file.name}")
        print("-" * 60)

        try:
            passages = extractor.extract_from_file(str(policy_file))
            formatted_passages = extractor.convert_to_compliance_format(passages)

            individual_output = output_path / f"{policy_file.stem}_for_mapping.json"
            with open(individual_output, "w", encoding="utf-8") as f:
                json.dump(formatted_passages, f, indent=2, ensure_ascii=False)

            print(f"  ✓ Saved {len(formatted_passages)} passages → {individual_output.name}")
            saved.append((policy_file.name, individual_output, len(formatted_passages)))

        except Exception as e:
            print(f"  ❌ Error: {e}")
            import traceback
            traceback.print_exc()
            continue

    if saved:
        print(f"\n{'='*60}")
        print("✅ COMPLETE!")
        print(f"{'='*60}")
        print(f"Processed {len(saved)} policy document(s)")
        for name, path, n in saved:
            print(f"   • {path.name} ({n} passages)")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Flexible policy extractor (use --backend docling or unstructured to preserve nested structure)"
    )
    parser.add_argument("--input-dir", type=str, default="data/01_raw/policies", help="Directory with policy files")
    parser.add_argument("--output-dir", type=str, default="data/02_processed/policies", help="Output directory")
    parser.add_argument(
        "--backend",
        type=str,
        default="pdfplumber",
        choices=["pdfplumber", "docling", "unstructured"],
        help="PDF parser: pdfplumber (default), docling (layout-aware), or unstructured (element-level)",
    )
    args = parser.parse_args()
    extract_all_policies_flexible(args.input_dir, args.output_dir, backend=args.backend)
