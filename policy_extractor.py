"""
Policy Document Extractor
Extracts and structures policy documents similar to UAE IA Regulation extraction

Processes policy PDFs and creates structured JSON format compatible with compliance mapping
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: pdfplumber not available. Install with: pip install pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")


@dataclass
class PolicySection:
    """Represents a section from a policy document"""
    policy_id: str  # e.g., "POL-001", "ACC-001"
    policy_name: str
    section_id: str  # e.g., "1.1", "4.2.1"
    section_title: str
    section_text: str
    subsection_texts: List[str]  # Sub-sections or bullet points
    page_number: int
    metadata: Dict


@dataclass
class StructuredPolicy:
    """Complete structured policy document"""
    policy_document_id: str
    policy_document_name: str
    version: str
    last_updated: str
    sections: List[PolicySection]
    metadata: Dict


class PolicyExtractor:
    """Extract and structure policy documents from PDF"""
    
    def __init__(self):
        # Common policy section patterns (order matters: more specific first)
        self.section_patterns = [
            r'^(\d+(?:\.\d+)+)\s+(.+?)$',  # 1.1 Section Title (requires at least one dot)
            r'^(\d+(?:\.\d+)*)\s+([A-Z][^.]{3,100})$',  # 1.1 TITLE (capitalized title)
            r'^Section\s+(\d+(?:\.\d+)*)\s*[:–-]\s*(.+?)$',  # Section 1.1: Title
            r'^(\d+\.\d+)\s+(.+?)$',  # 1.1 Title
            r'^(\d+)\)\s+(.+?)$',  # 1) Title
            r'^(\d+)\.\s+([A-Za-z].+?)$',  # 1. Title (number, period, space, title)
            r'^([A-Z]\d+)\s+(.+?)$',  # POL-001 Title
            r'^(\d+)\s+([A-Z][^.]{3,100})$',  # 1 TITLE (single number with capitalized title)
        ]
        # Standalone all-caps line (e.g. INTRODUCTION, PURPOSE, SCOPE) - single group, use for both id and title
        self.section_pattern_allcaps = re.compile(r'^([A-Z][A-Z0-9\s\-]{3,80})$')
        
        # Policy document identifiers
        self.policy_id_patterns = [
            r'Policy\s+(?:ID|Number|#)[:\s]+([A-Z0-9\-]+)',
            r'Policy\s+([A-Z0-9\-]+)',
            r'Document\s+ID[:\s]+([A-Z0-9\-]+)',
        ]
    
    def extract_from_pdf(self, pdf_path: str, policy_name: Optional[str] = None) -> StructuredPolicy:
        """
        Extract policy structure from PDF
        
        Args:
            pdf_path: Path to policy PDF
            policy_name: Optional policy name (extracted from PDF if not provided)
        """
        print(f"Extracting policy from: {pdf_path}")
        
        pdf_file = Path(pdf_path)
        policy_doc_id = pdf_file.stem  # Use filename as ID
        
        sections = []
        current_section = None
        current_section_text = []
        current_section_id = None
        current_section_title = None
        
        # Extract metadata
        metadata = self._extract_metadata(pdf_path)
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                if page_num % 10 == 0:
                    print(f"Processing page {page_num}/{len(pdf.pages)}...", end='\r')
                
                text = page.extract_text()
                if not text:
                    continue
                
                # Clean text
                text = self._clean_text(text)
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if this is a new section
                    section_match = self._match_section_pattern(line)
                    
                    if section_match:
                        # Save previous section
                        if current_section_id and current_section_text:
                            section_text = '\n'.join(current_section_text)
                            if len(section_text) > 20:  # Minimum length
                                sections.append(PolicySection(
                                    policy_id=policy_doc_id,
                                    policy_name=policy_name or policy_doc_id,
                                    section_id=current_section_id,
                                    section_title=current_section_title or '',
                                    section_text=section_text,
                                    subsection_texts=[],
                                    page_number=page_num - 1,  # Previous page
                                    metadata={'extracted_from': 'pdf'}
                                ))
                        
                        # Start new section
                        current_section_id = section_match[0]
                        current_section_title = section_match[1]
                        current_section_text = [line]
                    else:
                        # Accumulate text for current section
                        if current_section_id:
                            current_section_text.append(line)
                        else:
                            # Before first section, try to extract policy name
                            if not policy_name:
                                policy_name = self._extract_policy_name(line)
            
            # Don't forget the last section
            if current_section_id and current_section_text:
                section_text = '\n'.join(current_section_text)
                if len(section_text) > 20:
                    sections.append(PolicySection(
                        policy_id=policy_doc_id,
                        policy_name=policy_name or policy_doc_id,
                        section_id=current_section_id,
                        section_title=current_section_title or '',
                        section_text=section_text,
                        subsection_texts=[],
                        page_number=page_num,
                        metadata={'extracted_from': 'pdf'}
                    ))
        
        # If no sections found, create a single section with all text
        if len(sections) == 0:
            print("⚠️  No sections detected with patterns. Creating single section with all text...")
            # Re-extract all text
            all_text = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text.append(self._clean_text(text))
            
            full_text = '\n'.join(all_text)
            if len(full_text) > 50:
                sections.append(PolicySection(
                    policy_id=policy_doc_id,
                    policy_name=policy_name or policy_doc_id,
                    section_id="1",
                    section_title="Full Document",
                    section_text=full_text,
                    subsection_texts=[],
                    page_number=1,
                    metadata={'extracted_from': 'pdf', 'no_sections_detected': True}
                ))
        
        print(f"\n✓ Extracted {len(sections)} sections from PDF")
        
        return StructuredPolicy(
            policy_document_id=policy_doc_id,
            policy_document_name=policy_name or policy_doc_id,
            version=metadata.get('version', '1.0'),
            last_updated=metadata.get('last_updated', datetime.now().strftime('%Y-%m-%d')),
            sections=sections,
            metadata=metadata
        )
    
    def extract_from_docx(self, docx_path: str, policy_name: Optional[str] = None) -> StructuredPolicy:
        """
        Extract policy structure from DOCX
        
        Args:
            docx_path: Path to policy DOCX
            policy_name: Optional policy name (extracted from DOCX if not provided)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX extraction")
        
        print(f"Extracting policy from DOCX: {docx_path}")
        
        docx_file = Path(docx_path)
        policy_doc_id = docx_file.stem  # Use filename as ID
        
        sections = []
        current_section_text = []
        current_section_id = None
        current_section_title = None
        
        # Extract metadata
        metadata = self._extract_metadata_docx(docx_path)
        
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            line = para.text.strip()
            if not line:
                continue
            
            # Check if this is a new section
            section_match = self._match_section_pattern(line)
            
            if section_match:
                # Save previous section
                if current_section_id and current_section_text:
                    section_text = '\n'.join(current_section_text)
                    if len(section_text) > 20:  # Minimum length
                        sections.append(PolicySection(
                            policy_id=policy_doc_id,
                            policy_name=policy_name or policy_doc_id,
                            section_id=current_section_id,
                            section_title=current_section_title or '',
                            section_text=section_text,
                            subsection_texts=[],
                            page_number=0,  # DOCX doesn't have page numbers
                            metadata={'extracted_from': 'docx'}
                        ))
                
                # Start new section
                current_section_id = section_match[0]
                current_section_title = section_match[1]
                current_section_text = [line]
            else:
                # Accumulate text for current section
                if current_section_id:
                    current_section_text.append(line)
                else:
                    # Before first section, try to extract policy name
                    if not policy_name:
                        policy_name = self._extract_policy_name(line)
        
        # Don't forget the last section
        if current_section_id and current_section_text:
            section_text = '\n'.join(current_section_text)
            if len(section_text) > 20:
                sections.append(PolicySection(
                    policy_id=policy_doc_id,
                    policy_name=policy_name or policy_doc_id,
                    section_id=current_section_id,
                    section_title=current_section_title or '',
                    section_text=section_text,
                    subsection_texts=[],
                    page_number=0,
                    metadata={'extracted_from': 'docx'}
                ))
        
        # If no sections found, create a single section with all text
        if len(sections) == 0:
            print("⚠️  No sections detected with patterns. Creating single section with all text...")
            all_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text.append(text)
            
            full_text = '\n'.join(all_text)
            if len(full_text) > 50:
                sections.append(PolicySection(
                    policy_id=policy_doc_id,
                    policy_name=policy_name or policy_doc_id,
                    section_id="1",
                    section_title="Full Document",
                    section_text=full_text,
                    subsection_texts=[],
                    page_number=0,
                    metadata={'extracted_from': 'docx', 'no_sections_detected': True}
                ))
        
        print(f"\n✓ Extracted {len(sections)} sections from DOCX")
        
        return StructuredPolicy(
            policy_document_id=policy_doc_id,
            policy_document_name=policy_name or policy_doc_id,
            version=metadata.get('version', '1.0'),
            last_updated=metadata.get('last_updated', datetime.now().strftime('%Y-%m-%d')),
            sections=sections,
            metadata=metadata
        )
    
    def extract_from_file(self, file_path: str, policy_name: Optional[str] = None) -> StructuredPolicy:
        """
        Extract policy from file (auto-detects PDF or DOCX)
        
        Args:
            file_path: Path to policy file
            policy_name: Optional policy name
        """
        file_path_obj = Path(file_path)
        
        if file_path_obj.suffix.lower() == '.pdf':
            return self.extract_from_pdf(file_path, policy_name)
        elif file_path_obj.suffix.lower() in ['.docx', '.doc']:
            return self.extract_from_docx(file_path, policy_name)
        else:
            raise ValueError(f"Unsupported file format: {file_path_obj.suffix}")
    
    def _extract_metadata_docx(self, docx_path: str) -> Dict:
        """Extract metadata from DOCX"""
        metadata = {}
        
        try:
            doc = Document(docx_path)
            
            # Check first few paragraphs for metadata
            for para in doc.paragraphs[:10]:
                text = para.text
                
                # Look for version
                version_match = re.search(r'Version[:\s]+([\d.]+)', text, re.IGNORECASE)
                if version_match:
                    metadata['version'] = version_match.group(1)
                
                # Look for date
                date_match = re.search(r'(?:Date|Last\s+Updated|Effective)[:\s]+(\d{4}[-/]\d{2}[-/]\d{2})', text, re.IGNORECASE)
                if date_match:
                    metadata['last_updated'] = date_match.group(1)
                
                # Look for policy ID
                for pattern in self.policy_id_patterns:
                    id_match = re.search(pattern, text, re.IGNORECASE)
                    if id_match:
                        metadata['policy_id'] = id_match.group(1)
                        break
        except:
            pass
        
        return metadata
    
    def _match_section_pattern(self, line: str) -> Optional[Tuple[str, str]]:
        """Match line against section patterns"""
        for pattern in self.section_patterns:
            match = re.match(pattern, line)
            if match:
                section_id = match.group(1)
                section_title = match.group(2).strip()
                # Validate: section ID should be short, title should be meaningful
                # Also check that it's not just a number (like page numbers)
                is_valid = (
                    len(section_id) < 20 and 
                    len(section_title) > 3 and 
                    len(section_title) < 200 and
                    not section_title.lower().startswith('page') and
                    ('.' in section_id or not section_id.isdigit())  # Allow numbered sections with dots
                )
                if is_valid:
                    return (section_id, section_title)
        return None
    
    def _extract_policy_name(self, line: str) -> Optional[str]:
        """Extract policy name from line"""
        # Look for policy name patterns
        patterns = [
            r'Policy[:\s]+(.+?)(?:\n|$)',
            r'Document[:\s]+(.+?)(?:\n|$)',
            r'^(.+?)\s+Policy',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                if 5 < len(name) < 100:
                    return name
        
        return None
    
    def _extract_metadata(self, pdf_path: str) -> Dict:
        """Extract metadata from PDF (version, date, etc.)"""
        metadata = {}
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # Check first page for metadata
                first_page = pdf.pages[0]
                text = first_page.extract_text()
                
                if text:
                    # Look for version
                    version_match = re.search(r'Version[:\s]+([\d.]+)', text, re.IGNORECASE)
                    if version_match:
                        metadata['version'] = version_match.group(1)
                    
                    # Look for date
                    date_match = re.search(r'(?:Date|Last\s+Updated|Effective)[:\s]+(\d{4}[-/]\d{2}[-/]\d{2})', text, re.IGNORECASE)
                    if date_match:
                        metadata['last_updated'] = date_match.group(1)
                    
                    # Look for policy ID
                    for pattern in self.policy_id_patterns:
                        id_match = re.search(pattern, text, re.IGNORECASE)
                        if id_match:
                            metadata['policy_id'] = id_match.group(1)
                            break
        except:
            pass
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean PDF artifacts from text"""
        # Remove page numbers
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b\d+\s*\n\s*\n', '', text)
        
        # Remove headers/footers
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n', text)
        
        return text.strip()
    
    def convert_to_compliance_format(self, structured_policy: StructuredPolicy) -> List[Dict]:
        """
        Convert structured policy to format compatible with compliance mapping
        
        Returns list of policy passages in format expected by compliance_mapping_pipeline
        """
        policy_passages = []
        
        for section in structured_policy.sections:
            # Create passage for main section
            passage = {
                'id': f"{structured_policy.policy_document_id}_{section.section_id}",
                'name': f"{structured_policy.policy_document_name} - {section.section_title}",
                'text': section.section_text,
                'section': section.section_id,
                'metadata': {
                    'policy_id': structured_policy.policy_document_id,
                    'policy_name': structured_policy.policy_document_name,
                    'section_id': section.section_id,
                    'section_title': section.section_title,
                    'version': structured_policy.version,
                    'last_updated': structured_policy.last_updated,
                    'page_number': section.page_number
                }
            }
            policy_passages.append(passage)
            
            # Also create passages for subsections if they're substantial
            for idx, subsection in enumerate(section.subsection_texts):
                if len(subsection) > 50:  # Only substantial subsections
                    sub_passage = {
                        'id': f"{structured_policy.policy_document_id}_{section.section_id}.{idx+1}",
                        'name': f"{structured_policy.policy_document_name} - {section.section_title} (Subsection {idx+1})",
                        'text': subsection,
                        'section': f"{section.section_id}.{idx+1}",
                        'metadata': {
                            'policy_id': structured_policy.policy_document_id,
                            'is_subsection': True,
                            'parent_section': section.section_id
                        }
                    }
                    policy_passages.append(sub_passage)
        
        return policy_passages
    
    def save_policy(self, structured_policy: StructuredPolicy, output_path: str):
        """Save structured policy to JSON"""
        output_data = {
            'meta': {
                'policy_document_id': structured_policy.policy_document_id,
                'policy_document_name': structured_policy.policy_document_name,
                'version': structured_policy.version,
                'last_updated': structured_policy.last_updated,
                'total_sections': len(structured_policy.sections),
                'extracted_at': datetime.now().isoformat()
            },
            'sections': [asdict(section) for section in structured_policy.sections],
            'metadata': structured_policy.metadata
        }
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Saved structured policy to {output_path}")
    
    def save_for_compliance_mapping(self, structured_policy: StructuredPolicy, output_path: str):
        """Save policy in format ready for compliance mapping"""
        passages = self.convert_to_compliance_format(structured_policy)
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(passages, f, indent=2, ensure_ascii=False)
        
        print(f"✓ Saved {len(passages)} policy passages for compliance mapping to {output_path}")


def extract_all_policies(policies_dir: str = "data/01_raw/policies",
                         output_dir: str = "data/02_processed/policies"):
    """
    Extract all policy files (PDF and DOCX) from directory
    
    Args:
        policies_dir: Directory containing policy files
        output_dir: Directory to save extracted JSON files
    """
    policies_path = Path(policies_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Find all policy files
    pdf_files = list(policies_path.glob("*.pdf"))
    docx_files = list(policies_path.glob("*.docx"))
    all_files = pdf_files + docx_files
    
    if not all_files:
        print(f"⚠️  No policy files (PDF/DOCX) found in {policies_dir}")
        print(f"   Place your policy files in: {policies_dir}")
        return
    
    print(f"Found {len(pdf_files)} PDF(s) and {len(docx_files)} DOCX file(s)")
    print("=" * 60)
    
    extractor = PolicyExtractor()
    all_passages = []  # Collect all passages for combined file
    
    for policy_file in all_files:
        print(f"\nProcessing: {policy_file.name}")
        print("-" * 60)
        
        try:
            # Extract policy (auto-detects format)
            structured_policy = extractor.extract_from_file(str(policy_file))
            
            # Save structured version
            structured_output = output_path / f"{policy_file.stem}_structured.json"
            extractor.save_policy(structured_policy, str(structured_output))
            
            # Save for compliance mapping
            compliance_output = output_path / f"{policy_file.stem}_for_mapping.json"
            extractor.save_for_compliance_mapping(structured_policy, str(compliance_output))
            
            # Collect passages
            passages = extractor.convert_to_compliance_format(structured_policy)
            all_passages.extend(passages)
            
        except Exception as e:
            print(f"❌ Error processing {policy_file.name}: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    # Save combined file for compliance mapping
    if all_passages:
        combined_output = output_path / "all_policies_for_mapping.json"
        with open(combined_output, 'w', encoding='utf-8') as f:
            json.dump(all_passages, f, indent=2, ensure_ascii=False)
        
        print(f"\n{'='*60}")
        print(f"✅ COMPLETE!")
        print(f"{'='*60}")
        print(f"Processed {len(pdf_files)} policy document(s)")
        print(f"Extracted {len(all_passages)} policy passages")
        print(f"\n📁 Output files:")
        print(f"   - Individual structured: {output_path}/*_structured.json")
        print(f"   - Individual for mapping: {output_path}/*_for_mapping.json")
        print(f"   - Combined for mapping: {combined_output}")
        print(f"\n📋 Next step:")
        print(f"   Update quick_start_compliance.py to use:")
        print(f"   pipeline.load_policy_passages('{combined_output}')")


def main():
    """Main extraction function"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract and structure policy documents")
    parser.add_argument("--input-dir", type=str, default="data/01_raw/policies",
                       help="Directory containing policy PDFs")
    parser.add_argument("--output-dir", type=str, default="data/02_processed/policies",
                       help="Directory to save extracted JSON files")
    parser.add_argument("--single-file", type=str, help="Extract single PDF file")
    
    args = parser.parse_args()
    
    if args.single_file:
        # Extract single file (auto-detects format)
        extractor = PolicyExtractor()
        structured_policy = extractor.extract_from_file(args.single_file)
        
        output_path = Path(args.output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        pdf_name = Path(args.single_file).stem
        extractor.save_policy(structured_policy, str(output_path / f"{pdf_name}_structured.json"))
        extractor.save_for_compliance_mapping(structured_policy, 
                                            str(output_path / f"{pdf_name}_for_mapping.json"))
    else:
        # Extract all policy files (PDF and DOCX) in directory
        extract_all_policies(args.input_dir, args.output_dir)


if __name__ == "__main__":
    main()
