"""
Policy Mapping Extractor – RegNLP-aligned extraction for regulatory mapping.

Extracts from policy documents (e.g. Logging and Monitoring Policy) the elements
required to map to a "Source" regulation (e.g. UAE IA):

1. Document Metadata: Document ID, Title, Owner, Effective Dates
2. Hierarchical Passages: Sections by ToC (e.g. 5.1, 5.3, 5.7)
3. Actionable Obligations: Discrete "shall"/mandatory duties per section
4. Roles and Responsibilities: Responsibility matrix
5. Scope and Exceptions: Applicability and deviation process

Output is a structured JSON suitable as "Target" for mapping to UAE IA (Source).
"""

import re
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict, field


try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


# --- Output schema (mapping-ready) ---

@dataclass
class DocumentMetadata:
    document_id: str
    title: str
    owner: str = ""
    effective_date: str = ""
    version: str = ""
    classification: str = ""


@dataclass
class Obligation:
    text: str
    type: str  # "shall", "must", "required", etc.
    section_id: str = ""


@dataclass
class HierarchicalPassage:
    section_id: str       # e.g. "5.1", "5.3"
    section_title: str    # e.g. "Logging Requirements"
    content: str
    obligations: List[Obligation] = field(default_factory=list)


@dataclass
class RoleResponsibility:
    role: str
    responsibility: str
    source_section: str = ""


@dataclass
class PolicyMappingDocument:
    document_metadata: DocumentMetadata
    hierarchical_passages: List[HierarchicalPassage]
    roles_and_responsibilities: List[RoleResponsibility]
    scope: str = ""
    exceptions: str = ""
    source_path: str = ""


# --- Obligation detection ---

OBLIGATION_PATTERNS = [
    (re.compile(r"[^.]*\bshall\b[^.]*\.", re.IGNORECASE), "shall"),
    (re.compile(r"[^.]*\bmust\b[^.]*\.", re.IGNORECASE), "must"),
    (re.compile(r"[^.]*\bis required to\b[^.]*\.", re.IGNORECASE), "required"),
    (re.compile(r"[^.]*\bare required to\b[^.]*\.", re.IGNORECASE), "required"),
    (re.compile(r"[^.]*\bmandatory\b[^.]*\.", re.IGNORECASE), "mandatory"),
    (re.compile(r"[^.]*\bshall not\b[^.]*\.", re.IGNORECASE), "shall_not"),
    (re.compile(r"[^.]*\bmust not\b[^.]*\.", re.IGNORECASE), "must_not"),
]


def extract_obligations_from_text(text: str, section_id: str = "") -> List[Obligation]:
    """Identify discrete actionable obligations (shall, must, required, etc.)."""
    obligations = []
    for pattern, obl_type in OBLIGATION_PATTERNS:
        for m in pattern.finditer(text):
            sentence = m.group(0).strip()
            if len(sentence) < 15:
                continue
            obligations.append(Obligation(text=sentence, type=obl_type, section_id=section_id))
    return obligations


def normalize_section_id(heading: str) -> Tuple[str, str]:
    """
    Parse heading into section_id and title.
    e.g. "5.1 Logging Requirements" -> ("5.1", "Logging Requirements")
         "Logging Requirements" -> ("", "Logging Requirements")
    """
    heading = (heading or "").strip()
    match = re.match(r"^(\d+(?:\.\d+)*)\s+(.+)$", heading)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return "", heading


class PolicyMappingExtractor:
    """
    Extracts policy document elements for regulatory mapping (Target dataset).
    """

    def __init__(self):
        self._section_id_re = re.compile(r"^(\d+(?:\.\d+)*)\s+")

    def extract_from_docx(self, docx_path: str) -> PolicyMappingDocument:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required")
        path = Path(docx_path)
        doc = Document(docx_path)
        full_text = self._get_full_text(doc)

        # 1. Document metadata
        metadata = self._extract_document_metadata(doc, docx_path, full_text)

        # 2. Hierarchical passages (by numbered sections / ToC)
        passages = self._extract_hierarchical_passages(doc)

        # 3. Obligations are already attached per passage

        # 4. Roles and responsibilities
        roles = self._extract_roles_and_responsibilities(doc, full_text)

        # 5. Scope and exceptions
        scope = self._extract_scope(doc, full_text)
        exceptions = self._extract_exceptions(doc, full_text)

        return PolicyMappingDocument(
            document_metadata=metadata,
            hierarchical_passages=passages,
            roles_and_responsibilities=roles,
            scope=scope,
            exceptions=exceptions,
            source_path=str(path.resolve()),
        )

    def extract_from_pdf(self, pdf_path: str) -> PolicyMappingDocument:
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber is required")
        path = Path(pdf_path)
        full_text = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    full_text.append(t)
        full_text = "\n".join(full_text)
        metadata = self._extract_metadata_from_text(full_text, pdf_path)
        passages = self._extract_passages_from_pdf_text(full_text)
        roles = self._extract_roles_from_text(full_text)
        scope = self._extract_scope_from_text(full_text)
        exceptions = self._extract_exceptions_from_text(full_text)
        return PolicyMappingDocument(
            document_metadata=metadata,
            hierarchical_passages=passages,
            roles_and_responsibilities=roles,
            scope=scope,
            exceptions=exceptions,
            source_path=str(path.resolve()),
        )

    def extract_from_file(self, file_path: str) -> PolicyMappingDocument:
        p = Path(file_path)
        if p.suffix.lower() in (".docx", ".doc"):
            return self.extract_from_docx(file_path)
        if p.suffix.lower() == ".pdf":
            return self.extract_from_pdf(file_path)
        raise ValueError(f"Unsupported format: {p.suffix}")

    def _get_full_text(self, doc: "Document") -> str:
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _extract_document_metadata(
        self, doc: "Document", docx_path: str, full_text: str
    ) -> DocumentMetadata:
        """Extract Document ID, Title, Owner, Effective Dates from start of document or first table."""
        path = Path(docx_path)
        default_id = path.stem
        title = path.stem
        owner = ""
        effective_date = ""
        version = ""
        classification = ""

        # Try first table (common for policy cover/metadata)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                for i, cell in enumerate(cells):
                    cell_lower = cell.lower()
                    if "document id" in cell_lower or "policy id" in cell_lower:
                        if i + 1 < len(cells):
                            default_id = cells[i + 1] or default_id
                    elif "title" in cell_lower or "policy name" in cell_lower:
                        if i + 1 < len(cells):
                            title = cells[i + 1] or title
                    elif "owner" in cell_lower or "approver" in cell_lower:
                        if i + 1 < len(cells):
                            owner = cells[i + 1] or owner
                    elif "effective" in cell_lower or "date" in cell_lower:
                        if i + 1 < len(cells):
                            effective_date = cells[i + 1] or effective_date
                    elif "version" in cell_lower:
                        if i + 1 < len(cells):
                            version = cells[i + 1] or version
                    elif "classification" in cell_lower:
                        if i + 1 < len(cells):
                            classification = cells[i + 1] or classification
            break  # first table only for metadata

        # Fallback: try first few paragraphs for title
        for para in doc.paragraphs[:5]:
            t = para.text.strip()
            if len(t) > 5 and len(t) < 200 and not re.match(r"^\d+\.", t):
                if not title or title == default_id:
                    title = t
                break

        return DocumentMetadata(
            document_id=default_id,
            title=title,
            owner=owner,
            effective_date=effective_date,
            version=version,
            classification=classification,
        )

    def _extract_metadata_from_text(self, text: str, source_path: str) -> DocumentMetadata:
        path = Path(source_path)
        default_id = path.stem
        title = default_id
        for line in text.split("\n")[:30]:
            line = line.strip()
            if re.match(r"Document\s+ID|Policy\s+ID", line, re.IGNORECASE):
                parts = re.split(r"[:–\-]\s*", line, maxsplit=1)
                if len(parts) > 1:
                    default_id = parts[1].strip()
            if re.match(r"Title|Policy\s+Name", line, re.IGNORECASE):
                parts = re.split(r"[:–\-]\s*", line, maxsplit=1)
                if len(parts) > 1:
                    title = parts[1].strip()
        return DocumentMetadata(document_id=default_id, title=title)

    def _extract_hierarchical_passages(self, doc: "Document") -> List[HierarchicalPassage]:
        """Break document into granular sections based on numbered headings (ToC structure)."""
        passages = []
        current_section_id = ""
        current_title = ""
        current_paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_numbered_heading = self._is_numbered_heading(para, text)
            if is_numbered_heading:
                if current_paragraphs:
                    content = " ".join(current_paragraphs)
                    if content.strip():
                        obligations = extract_obligations_from_text(content, current_section_id)
                        passages.append(
                            HierarchicalPassage(
                                section_id=current_section_id,
                                section_title=current_title,
                                content=content,
                                obligations=obligations,
                            )
                        )
                current_section_id, current_title = normalize_section_id(text)
                current_paragraphs = []
            else:
                current_paragraphs.append(text)

        if current_paragraphs:
            content = " ".join(current_paragraphs)
            if content.strip():
                obligations = extract_obligations_from_text(content, current_section_id)
                passages.append(
                    HierarchicalPassage(
                        section_id=current_section_id,
                        section_title=current_title,
                        content=content,
                        obligations=obligations,
                    )
                )
        return passages

    def _is_numbered_heading(self, para: "Paragraph", text: str) -> bool:
        if self._section_id_re.match(text):
            return True
        style = (para.style and para.style.name or "").lower()
        if "heading" in style or "title" in style:
            return True
        if para.runs and any(r.bold for r in para.runs) and len(text) < 120:
            return True
        return False

    def _extract_passages_from_pdf_text(self, text: str) -> List[HierarchicalPassage]:
        passages = []
        current_id, current_title = "", ""
        current_lines = []
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            sid, stitle = normalize_section_id(line)
            if sid or (len(line) < 100 and re.match(r"^\d+[\.\)]", line)):
                if current_lines:
                    content = " ".join(current_lines)
                    if content.strip():
                        obligations = extract_obligations_from_text(content, current_id)
                        passages.append(
                            HierarchicalPassage(
                                section_id=current_id,
                                section_title=current_title or "Section",
                                content=content,
                                obligations=obligations,
                            )
                        )
                current_id, current_title = sid or current_id, stitle or line
                current_lines = []
            else:
                current_lines.append(line)
        if current_lines:
            content = " ".join(current_lines)
            if content.strip():
                obligations = extract_obligations_from_text(content, current_id)
                passages.append(
                    HierarchicalPassage(
                        section_id=current_id,
                        section_title=current_title or "Section",
                        content=content,
                        obligations=obligations,
                    )
                )
        return passages

    def _extract_roles_and_responsibilities(
        self, doc: "Document", full_text: str
    ) -> List[RoleResponsibility]:
        """Extract responsibility matrix (section or table)."""
        roles = []
        in_section = False
        section_name = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if re.search(
                r"Roles?\s+and\s+Responsibilities|Responsibility\s+Matrix|R\s*&\s*R",
                text,
                re.IGNORECASE,
            ):
                in_section = True
                section_name = text
                continue
            if in_section:
                if self._section_id_re.match(text) and "Roles" not in text:
                    in_section = False
                    continue
                if re.match(r"^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s*[:–\-]", text):
                    parts = re.split(r"\s*[:–\-]\s*", text, maxsplit=1)
                    if len(parts) == 2:
                        roles.append(
                            RoleResponsibility(
                                role=parts[0].strip(),
                                responsibility=parts[1].strip(),
                                source_section=section_name,
                            )
                        )
                        continue
                if "\t" in text or "  " in text:
                    parts = re.split(r"\t|\s{2,}", text, maxsplit=1)
                    if len(parts) == 2 and len(parts[0]) < 80:
                        roles.append(
                            RoleResponsibility(
                                role=parts[0].strip(),
                                responsibility=parts[1].strip(),
                                source_section=section_name,
                            )
                        )

        for table in doc.tables:
            rows = [[c.text.strip() for c in row.cells] for row in table.rows]
            if not rows:
                continue
            header = " ".join(rows[0]).lower()
            if "role" not in header and "responsib" not in header:
                continue
            role_col, resp_col = 0, 1
            for i, cell in enumerate(rows[0]):
                if "role" in cell.lower():
                    role_col = i
                if "responsib" in cell.lower() or "duty" in cell.lower():
                    resp_col = i
            for row in rows[1:]:
                if len(row) > max(role_col, resp_col):
                    role = row[role_col] or ""
                    resp = row[resp_col] or ""
                    if role and resp:
                        roles.append(
                            RoleResponsibility(
                                role=role,
                                responsibility=resp,
                                source_section="Roles and Responsibilities (table)",
                            )
                        )
        return roles

    def _extract_roles_from_text(self, text: str) -> List[RoleResponsibility]:
        roles = []
        in_roles = False
        for line in text.split("\n"):
            line = line.strip()
            if re.search(r"Roles?\s+and\s+Responsibilities|Responsibility\s+Matrix", line, re.IGNORECASE):
                in_roles = True
                continue
            if in_roles and re.match(r"^\d+\.\d+", line):
                in_roles = False
                continue
            if in_roles and re.match(r"^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s*[:–\-]", line):
                parts = re.split(r"\s*[:–\-]\s*", line, maxsplit=1)
                if len(parts) == 2 and len(parts[1]) > 10:
                    roles.append(
                        RoleResponsibility(role=parts[0].strip(), responsibility=parts[1].strip())
                    )
        return roles

    def _extract_scope(self, doc: "Document", full_text: str) -> str:
        return self._extract_named_section(doc, full_text, ["Scope", "Applicability"])

    def _extract_scope_from_text(self, text: str) -> str:
        return self._extract_named_section_from_text(text, ["Scope", "Applicability"])

    def _extract_exceptions(self, doc: "Document", full_text: str) -> str:
        return self._extract_named_section(
            doc, full_text, ["Exception", "Exceptions", "Policy Exception", "Deviation"]
        )

    def _extract_exceptions_from_text(self, text: str) -> str:
        return self._extract_named_section_from_text(
            text, ["Exception", "Exceptions", "Policy Exception", "Deviation"]
        )

    def _extract_named_section(
        self, doc: "Document", full_text: str, section_titles: List[str]
    ) -> str:
        """Collect text from the first matching section until next numbered heading."""
        collecting = False
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if collecting:
                if self._section_id_re.match(text) and len(text) < 120:
                    return " ".join(paragraphs).strip()
                paragraphs.append(text)
                continue
            for title in section_titles:
                if re.match(re.escape(title) + r"\s*$", text, re.IGNORECASE):
                    collecting = True
                    paragraphs = []
                    break
        return " ".join(paragraphs).strip() if paragraphs else ""

    def _extract_named_section_from_text(self, text: str, section_titles: List[str]) -> str:
        lines = text.split("\n")
        collecting = False
        buf = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if collecting:
                if re.match(r"^\d+\.\d+", line) and len(line) < 120:
                    return " ".join(buf).strip()
                buf.append(line)
                continue
            for title in section_titles:
                if re.match(re.escape(title) + r"\s*$", line, re.IGNORECASE):
                    collecting = True
                    buf = []
                    break
        return " ".join(buf).strip() if buf else ""


def save_mapping_document(doc: PolicyMappingDocument, output_path: str) -> None:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    out = {
        "document_metadata": asdict(doc.document_metadata),
        "hierarchical_passages": [
            {
                "section_id": p.section_id,
                "section_title": p.section_title,
                "content": p.content,
                "obligations": [asdict(o) for o in p.obligations],
            }
            for p in doc.hierarchical_passages
        ],
        "roles_and_responsibilities": [asdict(r) for r in doc.roles_and_responsibilities],
        "scope": doc.scope,
        "exceptions": doc.exceptions,
        "source_path": doc.source_path,
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(out, f, indent=2, ensure_ascii=False)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Extract policy elements for regulatory mapping")
    parser.add_argument("input", nargs="?", default="data/01_raw/policies", help="Policy file or directory")
    parser.add_argument("--output-dir", default="data/02_processed/policies_mapping", help="Output directory")
    args = parser.parse_args()
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    input_path = Path(args.input)
    extractor = PolicyMappingExtractor()
    if input_path.is_file():
        files = [input_path]
    else:
        files = list(input_path.glob("*.docx")) + list(input_path.glob("*.pdf"))
    if not files:
        print("No .docx or .pdf files found.")
        return
    for f in files:
        print(f"Extracting: {f.name}")
        try:
            doc = extractor.extract_from_file(str(f))
            out_path = out_dir / f"{f.stem}_mapping.json"
            save_mapping_document(doc, str(out_path))
            print(f"  -> {len(doc.hierarchical_passages)} passages, "
                  f"{sum(len(p.obligations) for p in doc.hierarchical_passages)} obligations, "
                  f"{len(doc.roles_and_responsibilities)} roles")
            print(f"  Saved: {out_path}")
        except Exception as e:
            print(f"  Error: {e}")
            raise


if __name__ == "__main__":
    main()
